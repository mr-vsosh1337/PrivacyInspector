import os
import re
from typing import List

import fitz
from docx import Document
import openpyxl

from presidio_analyzer import AnalyzerEngine

from .redact import redact_text, get_redaction_results


def make_out_path(in_path: str, out_dir: str) -> str:
    base = os.path.basename(in_path)
    name, ext = os.path.splitext(base)
    return os.path.join(out_dir, f"{name}.redacted{ext}")


def redact_txt(in_path: str, out_path: str, analyzer: AnalyzerEngine, token: str,
               excluded_texts: set[str] | None = None) -> None:
    with open(in_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    redacted = redact_text(analyzer, text, token, excluded_texts)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(redacted)


def _apply_text_to_runs(paragraph, new_text: str) -> None:
    if not paragraph.runs:
        paragraph.text = new_text
        return

    start = 0
    last_index = len(paragraph.runs) - 1
    for index, run in enumerate(paragraph.runs):
        if index == last_index:
            run.text = new_text[start:]
            break
        run_len = len(run.text)
        run.text = new_text[start:start + run_len]
        start += run_len


def _redact_paragraph(paragraph, analyzer: AnalyzerEngine, token: str, excluded_texts: set[str] | None = None) -> None:
    original_text = paragraph.text
    if not original_text:
        return

    new_text = redact_text(analyzer, original_text, token, excluded_texts)
    if new_text == original_text:
        return

    _apply_text_to_runs(paragraph, new_text)


def _iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested_table in cell.tables:
                yield from _iter_table_paragraphs(nested_table)


def _iter_header_footer_paragraphs(doc):
    for section in doc.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                yield paragraph
            for table in container.tables:
                yield from _iter_table_paragraphs(table)


def redact_docx(in_path: str, out_path: str, analyzer: AnalyzerEngine, token: str,
                excluded_texts: set[str] | None = None) -> None:
    doc = Document(in_path)

    for paragraph in doc.paragraphs:
        _redact_paragraph(paragraph, analyzer, token, excluded_texts)

    for table in doc.tables:
        for paragraph in _iter_table_paragraphs(table):
            _redact_paragraph(paragraph, analyzer, token, excluded_texts)

    for paragraph in _iter_header_footer_paragraphs(doc):
        _redact_paragraph(paragraph, analyzer, token, excluded_texts)

    doc.save(out_path)


def redact_xlsx(in_path: str, out_path: str, analyzer: AnalyzerEngine, token: str,
                excluded_texts: set[str] | None = None) -> None:
    wb = openpyxl.load_workbook(in_path)
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and cell.value.strip():
                    new_val = redact_text(analyzer, cell.value, token, excluded_texts)
                    if new_val != cell.value:
                        cell.value = new_val
    wb.save(out_path)


def redact_pdf(in_path: str, out_path: str, analyzer: AnalyzerEngine, excluded_texts: set[str] | None = None) -> None:
    pdf = fitz.open(in_path)

    for page in pdf:
        page_text = page.get_text("text") or ""
        results = get_redaction_results(analyzer, page_text, excluded_texts)
        if not results:
            continue

        snippets: List[str] = []
        for r in results:
            snip = page_text[r.start:r.end].strip()
            snip = re.sub(r"\s+", " ", snip)
            if not snip:
                continue
            if len(snip) > 80:
                snip = snip[:80].strip()
            snippets.append(snip)

        def _norm(s: str) -> str:
            return re.sub(r"\s+", " ", s).strip()

        excluded_norm = {_norm(x) for x in (excluded_texts or set())}

        seen = set()
        snippets_unique: List[str] = []
        for s in snippets:
            ns = _norm(s)
            if ns in excluded_norm:
                continue
            if ns not in seen:
                seen.add(ns)
                snippets_unique.append(s)

        for snip in snippets_unique:
            rects = page.search_for(snip)
            for rect in rects:
                page.add_redact_annot(rect, fill=(0, 0, 0))

        page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_REMOVE)

    pdf.save(out_path, garbage=4, deflate=True)
    pdf.close()
