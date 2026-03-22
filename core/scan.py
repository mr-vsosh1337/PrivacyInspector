import os
import re
from collections import Counter, defaultdict
from dataclasses import dataclass
from typing import List

import fitz
from docx import Document
import openpyxl

from presidio_analyzer import AnalyzerEngine
from config import SUPPORTED_EXTS

from .redact import analyze_multilang, filter_label_spans


@dataclass(frozen=True)
class Finding:
    entity_type: str
    text: str
    count: int
    max_score: float


def _norm(s: str) -> str:
    return re.sub(r"\s+", " ", s).strip()


def extract_text(in_path: str) -> str:
    ext = os.path.splitext(in_path)[1].lower()

    if ext == ".txt":
        with open(in_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    if ext == ".docx":
        doc = Document(in_path)
        parts = []
        for p in doc.paragraphs:
            if p.text:
                parts.append(p.text)
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        return "\n".join(parts)

    if ext == ".xlsx":
        wb = openpyxl.load_workbook(in_path)
        parts = []
        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    if isinstance(cell.value, str) and cell.value.strip():
                        parts.append(cell.value)
        return "\n".join(parts)

    if ext == ".pdf":
        pdf = fitz.open(in_path)
        parts = []
        for page in pdf:
            parts.append(page.get_text("text") or "")
        pdf.close()
        return "\n".join(parts)

    return ""


def scan_file(analyzer: AnalyzerEngine, in_path: str) -> List[Finding]:
    text = extract_text(in_path)
    if not text.strip():
        return []

    results = analyze_multilang(analyzer, text)
    results = filter_label_spans(text, results)
    if not results:
        return []

    counts = Counter()
    max_scores = defaultdict(float)

    for r in results:
        span = _norm(text[r.start:r.end])
        if not span:
            continue
        key = (r.entity_type, span)
        counts[key] += 1
        if r.score > max_scores[key]:
            max_scores[key] = r.score

    findings = [
        Finding(entity_type=k[0], text=k[1], count=counts[k], max_score=max_scores[k])
        for k in counts
    ]

    findings.sort(key=lambda f: (-f.count, -f.max_score, f.entity_type, f.text))
    return findings
